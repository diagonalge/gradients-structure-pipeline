from __future__ import annotations

import gc
import hashlib
import json
import math
import os
import random
import subprocess
import sys
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from collections import defaultdict
from collections.abc import Iterator
from contextlib import contextmanager
from contextvars import ContextVar
from csv import DictReader
from dataclasses import dataclass
from itertools import islice
from pathlib import Path
from typing import Any

from datasets import load_dataset
from loguru import logger

from .models import SourceDocument


@dataclass(frozen=True)
class DatasetPreset:
    dataset_id: str
    text_column: str
    id_column: str | None = None
    title_column: str | None = None
    config: str | None = None


@dataclass(frozen=True)
class DiscoveredSchema:
    source: str
    source_kind: str
    text_column: str
    id_column: str | None
    title_column: str | None
    metadata_columns: tuple[str, ...]
    sampled_rows: int
    confidence: float
    text_candidates: tuple[dict[str, Any], ...] = ()

    def to_dict(self) -> dict[str, Any]:
        return {
            "source": self.source,
            "source_kind": self.source_kind,
            "text_column": self.text_column,
            "id_column": self.id_column,
            "title_column": self.title_column,
            "metadata_columns": list(self.metadata_columns),
            "sampled_rows": self.sampled_rows,
            "confidence": self.confidence,
            "text_candidates": list(self.text_candidates),
        }

    @classmethod
    def from_dict(cls, value: dict[str, Any]) -> DiscoveredSchema:
        return cls(
            source=value["source"],
            source_kind=value["source_kind"],
            text_column=value["text_column"],
            id_column=value.get("id_column"),
            title_column=value.get("title_column"),
            metadata_columns=tuple(value.get("metadata_columns", [])),
            sampled_rows=value["sampled_rows"],
            confidence=value["confidence"],
            text_candidates=tuple(value.get("text_candidates", [])),
        )


PRESETS: dict[str, DatasetPreset] = {
    "open-patients": DatasetPreset(
        dataset_id="abuhoraira06/Open-Patients",
        text_column="description",
        id_column="_id",
    ),
    "technetium": DatasetPreset(
        dataset_id="temlm-foundation/Technetium-I",
        text_column="text",
        id_column="note_id",
        title_column="note_type",
    ),
    "canadian-caselaw": DatasetPreset(
        dataset_id="a2aj/canadian-case-law",
        text_column="unofficial_text_en",
        id_column="citation_en",
        title_column="name_en",
    ),
    "bailii": DatasetPreset(
        dataset_id="easymn/bailii_260505_raw",
        text_column="html_content",
        id_column="path",
        title_column="title",
    ),
    "cold-cases": DatasetPreset(
        dataset_id="harvard-lil/cold-cases",
        text_column="opinions[].opinion_text",
        id_column="id",
        title_column="case_name",
    ),
}


def nested_values(row: dict[str, Any], path: str | None) -> list[Any]:
    if not path:
        return []
    values: list[Any] = [row]
    for raw_part in path.split("."):
        part = raw_part.removesuffix("[]")
        next_values: list[Any] = []
        for value in values:
            if isinstance(value, list):
                for item in value:
                    if isinstance(item, dict) and part in item:
                        next_values.append(item[part])
            elif isinstance(value, dict) and part in value:
                next_values.append(value[part])
        values = [item for value in next_values for item in value] if raw_part.endswith("[]") else next_values
    return values


def nested_value(row: dict[str, Any], path: str | None) -> Any:
    values = nested_values(row, path)
    return values[0] if values else None


def _joined_text(row: dict[str, Any], path: str) -> str:
    values = [value.strip() for value in nested_values(row, path) if isinstance(value, str) and value.strip()]
    return "\n\n".join(dict.fromkeys(values))


def _flatten_scalars(value: Any, path: str = "", depth: int = 0) -> Iterator[tuple[str, Any]]:
    if depth > 8:
        return
    if isinstance(value, dict):
        for key, child in value.items():
            child_path = f"{path}.{key}" if path else str(key)
            yield from _flatten_scalars(child, child_path, depth + 1)
    elif isinstance(value, list):
        list_path = f"{path}[]"
        for child in value[:8]:
            yield from _flatten_scalars(child, list_path, depth + 1)
    elif value is not None and path:
        yield path, value


def _path_tokens(path: str) -> set[str]:
    normalized = path.casefold().replace("[]", "").replace("-", "_")
    return {part for component in normalized.split(".") for part in component.split("_") if part}


def _text_name_score(path: str) -> float:
    leaf = path.casefold().replace("[]", "").split(".")[-1]
    exact = {
        "opinion_text": 14,
        "full_text": 14,
        "document_text": 14,
        "html_content": 13,
        "unofficial_text_en": 13,
        "text": 12,
        "content": 11,
        "body": 10,
        "document": 10,
        "description": 8,
        "opinion": 8,
        "transcript": 8,
        "article": 7,
        "summary": 4,
        "syllabus": 4,
        "headnotes": 3,
    }
    score = exact.get(leaf, 0)
    tokens = _path_tokens(path)
    if tokens & {"id", "uuid", "slug", "url", "path", "date", "name", "title", "citation"}:
        score -= 9
    return score


def _rank_text_candidates(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    stats: dict[str, list[str]] = defaultdict(list)
    for row in rows:
        for path, value in _flatten_scalars(row):
            if isinstance(value, str) and value.strip():
                stats[path].append(value.strip())

    candidates = []
    for path, values in stats.items():
        average_chars = sum(len(value) for value in values) / len(values)
        coverage = len(values) / max(len(rows), 1)
        score = _text_name_score(path) + min(math.log10(average_chars + 1) * 2.5, 11) + coverage * 4
        if average_chars < 80:
            score -= 8
        candidates.append(
            {
                "path": path,
                "score": round(score, 3),
                "coverage": round(min(coverage, 1.0), 3),
                "average_chars": round(average_chars),
            }
        )
    return sorted(candidates, key=lambda candidate: (candidate["score"], candidate["average_chars"]), reverse=True)


def _rank_scalar_path(
    rows: list[dict[str, Any]],
    *,
    preferred: tuple[str, ...],
    maximum_average_chars: int,
) -> str | None:
    stats: dict[str, list[Any]] = defaultdict(list)
    for row in rows:
        for path, value in _flatten_scalars(row):
            if "[]" not in path and isinstance(value, (str, int)) and value not in ("", None):
                stats[path].append(value)

    ranked: list[tuple[float, str]] = []
    for path, values in stats.items():
        average_chars = sum(len(str(value)) for value in values) / len(values)
        if average_chars > maximum_average_chars:
            continue
        leaf = path.casefold().split(".")[-1]
        name_score = max((12 - index * 1.5 for index, name in enumerate(preferred) if leaf == name), default=0)
        uniqueness = len({str(value) for value in values}) / len(values)
        coverage = len(values) / max(len(rows), 1)
        ranked.append((name_score + uniqueness * 4 + coverage * 2, path))
    return max(ranked, default=(0, None))[1]


def discover_schema(
    rows: list[dict[str, Any]],
    *,
    source: str,
    source_kind: str,
    text_column: str | None = None,
    id_column: str | None = None,
    title_column: str | None = None,
    metadata_columns: tuple[str, ...] = (),
) -> DiscoveredSchema:
    if not rows:
        raise ValueError("Cannot discover a schema without sample rows")
    text_candidates = _rank_text_candidates(rows)
    selected_text = text_column or (text_candidates[0]["path"] if text_candidates else None)
    if not selected_text:
        raise ValueError("No non-empty text-like field was discovered")
    selected_id = id_column or _rank_scalar_path(
        rows,
        preferred=("id", "_id", "uuid", "slug", "document_id", "case_id", "citation"),
        maximum_average_chars=300,
    )
    selected_title = title_column or _rank_scalar_path(
        rows,
        preferred=("title", "case_name", "name", "case_name_short", "case_name_full", "headline", "subject", "slug"),
        maximum_average_chars=500,
    )
    if metadata_columns:
        selected_metadata = metadata_columns
    else:
        excluded = {selected_text, selected_id, selected_title}
        metadata_candidates = [
            path
            for path in (
                "date_filed",
                "court_full_name",
                "court_short_name",
                "court_jurisdiction",
                "precedential_status",
                "disposition",
                "nature_of_suit",
                "language",
                "category",
                "jurisdiction",
                "date",
            )
            if path not in excluded and any(nested_values(row, path) for row in rows)
        ]
        selected_metadata = tuple(metadata_candidates[:8])
    margin = (
        text_candidates[0]["score"] - text_candidates[1]["score"]
        if len(text_candidates) > 1
        else text_candidates[0]["score"]
        if text_candidates
        else 0
    )
    confidence = 1.0 if text_column else max(0.05, min(0.99, 0.55 + margin / 20))
    return DiscoveredSchema(
        source=source,
        source_kind=source_kind,
        text_column=selected_text,
        id_column=selected_id,
        title_column=selected_title,
        metadata_columns=selected_metadata,
        sampled_rows=len(rows),
        confidence=round(confidence, 3),
        text_candidates=tuple(text_candidates[:5]),
    )


def inspect_hf_schema(
    dataset_id: str,
    *,
    split: str = "train",
    config: str | None = None,
    sample_rows: int = 32,
    text_column: str | None = None,
    id_column: str | None = None,
    title_column: str | None = None,
    metadata_columns: tuple[str, ...] = (),
) -> DiscoveredSchema:
    dataset = load_dataset(dataset_id, config, split=split, streaming=True)
    iterator = iter(dataset)
    try:
        rows = list(islice(iterator, sample_rows))
    finally:
        close = getattr(iterator, "close", None)
        if callable(close):
            close()
        del iterator
        del dataset
        gc.collect()
    return discover_schema(
        rows,
        source=dataset_id,
        source_kind="huggingface",
        text_column=text_column,
        id_column=id_column,
        title_column=title_column,
        metadata_columns=metadata_columns,
    )


def load_hf_documents_isolated(
    dataset_id: str,
    *,
    split: str = "train",
    config: str | None = None,
    limit: int,
    seed: int = 42,
    shuffle_buffer: int = 1_000,
    schema_sample_rows: int = 32,
    text_column: str | None = None,
    id_column: str | None = None,
    title_column: str | None = None,
    metadata_columns: tuple[str, ...] = (),
) -> tuple[list[SourceDocument], DiscoveredSchema]:
    with tempfile.TemporaryDirectory(prefix="unstructured-source-") as temporary_directory:
        payload_path = Path(temporary_directory) / "source.jsonl"
        command = [
            sys.executable,
            "-m",
            "ds_structure_server.pipeline.hf_source_worker",
            "--dataset",
            dataset_id,
            "--split",
            split,
            "--limit",
            str(limit),
            "--seed",
            str(seed),
            "--shuffle-buffer",
            str(shuffle_buffer),
            "--schema-sample-rows",
            str(schema_sample_rows),
            "--output",
            str(payload_path),
        ]
        for flag, value in (
            ("--config", config),
            ("--text-column", text_column),
            ("--id-column", id_column),
            ("--title-column", title_column),
        ):
            if value:
                command.extend((flag, value))
        for column in metadata_columns:
            command.extend(("--metadata-column", column))
        subprocess.run(command, check=True)

        lines_iter = payload_path.open("r", encoding="utf-8")
        try:
            first = next(lines_iter, None)
            if not first:
                raise ValueError("Hugging Face source worker returned no schema or documents")
            schema = DiscoveredSchema.from_dict(json.loads(first)["schema"])
            documents: list[SourceDocument] = []
            for line in lines_iter:
                text = line.strip()
                if not text:
                    continue
                value = json.loads(text)
                documents.append(
                    SourceDocument(
                        doc_id=value["doc_id"],
                        text=_truncate_chars(value["text"]),
                        title=value["title"],
                        metadata=value.get("metadata", {}),
                    )
                )
            return documents, schema
        finally:
            lines_iter.close()


_DOCUMENT_SUFFIXES = {".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx"}
_STRUCTURED_SUFFIXES = {".json", ".jsonl", ".csv"}
_ARCHIVE_SUFFIXES = {".zip"}
_MIN_PDF_CHARS_PER_PAGE = 40
_MIN_PDF_NONEMPTY_PAGE_RATIO = 0.2
# Defaults sized for a ~16 vCPU / 16 GiB structure worker.
# 0 for page/char caps means "no limit" (process the full document).
_DEFAULT_OCR_DPI = int(os.getenv("STRUCTURE_OCR_DPI", "100"))
_DEFAULT_OCR_MAX_PAGES = int(os.getenv("STRUCTURE_OCR_MAX_PAGES", "0"))
_DEFAULT_OCR_WORKERS = max(1, int(os.getenv("STRUCTURE_OCR_WORKERS", "16")))
# Keep each tesseract process single-threaded so STRUCTURE_OCR_WORKERS can saturate cores.
_OCR_OMP_THREADS = max(1, int(os.getenv("STRUCTURE_OCR_OMP_THREADS", "1")))
_DEFAULT_NATIVE_PDF_MAX_PAGES = int(os.getenv("STRUCTURE_PDF_MAX_PAGES", "0"))
_MAX_DOCUMENT_CHARS = int(os.getenv("STRUCTURE_MAX_DOCUMENT_CHARS", "0"))
_MAX_STRUCTURED_FILE_BYTES = int(os.getenv("STRUCTURE_MAX_STRUCTURED_FILE_BYTES", str(128 * 1024 * 1024)))
_MAX_ZIP_UNCOMPRESSED_BYTES = int(os.getenv("STRUCTURE_MAX_ZIP_UNCOMPRESSED_BYTES", str(512 * 1024 * 1024)))
# Suggest OCR skim for personas; capacity_chars is extrapolated to the full PDF.
# Pages are sampled from the document interior (not prefix) to avoid front-matter bias.
_SUGGEST_OCR_DPI = int(os.getenv("STRUCTURE_SUGGEST_OCR_DPI", "100"))
_SUGGEST_OCR_MAX_PAGES = int(os.getenv("STRUCTURE_SUGGEST_OCR_MAX_PAGES", "32"))
_SUGGEST_MAX_DOCUMENT_CHARS = int(os.getenv("STRUCTURE_SUGGEST_MAX_DOCUMENT_CHARS", "0"))
_LIGHT_DOCUMENT_MODE: ContextVar[bool] = ContextVar("structure_light_document_mode", default=False)


def _sample_page_numbers(
    total_pages: int,
    limit: int,
    *,
    seed: int | None = None,
) -> list[int]:
    """Pick up to ``limit`` 1-based pages, preferring interior over front/back matter."""
    if total_pages <= 0:
        return []
    if limit <= 0 or limit >= total_pages:
        return list(range(1, total_pages + 1))
    rng = random.Random(seed if seed is not None else (total_pages * 1_000_003 + limit))
    # Skip typical ProQuest / TOC / approval-sheet prefix and sparse trailing pages.
    skip_front = min(max(5, total_pages // 10), max(0, total_pages - limit))
    skip_back = min(max(2, total_pages // 20), max(0, total_pages - skip_front - limit))
    lo = 1 + skip_front
    hi = total_pages - skip_back
    if hi < lo:
        lo, hi = 1, total_pages
    pool = list(range(lo, hi + 1))
    if len(pool) >= limit:
        return sorted(rng.sample(pool, limit))
    chosen = set(pool)
    remainder = [p for p in range(1, total_pages + 1) if p not in chosen]
    need = limit - len(chosen)
    if need > 0 and remainder:
        chosen.update(rng.sample(remainder, min(need, len(remainder))))
    return sorted(chosen)


@contextmanager
def light_document_mode(enabled: bool = True):
    """Bound PDF/OCR/text extraction for suggest/discovery (not full generation)."""
    token = _LIGHT_DOCUMENT_MODE.set(bool(enabled))
    try:
        yield
    finally:
        _LIGHT_DOCUMENT_MODE.reset(token)


def _in_light_document_mode() -> bool:
    return bool(_LIGHT_DOCUMENT_MODE.get())


def _safe_extract_zip(zip_path: Path, destination: Path) -> Path:
    import shutil
    import zipfile

    destination.mkdir(parents=True, exist_ok=True)
    dest_root = destination.resolve()
    uncompressed = 0
    with zipfile.ZipFile(zip_path, "r") as archive:
        for info in archive.infolist():
            if info.is_dir():
                continue
            member = Path(info.filename)
            if member.is_absolute() or ".." in member.parts:
                continue
            size = int(info.file_size or 0)
            uncompressed += size
            if uncompressed > _MAX_ZIP_UNCOMPRESSED_BYTES:
                raise ValueError(
                    f"Zip uncompressed size exceeds {_MAX_ZIP_UNCOMPRESSED_BYTES // (1024 * 1024)}MB: {zip_path}"
                )
            target = (destination / member).resolve()
            if not str(target).startswith(str(dest_root)):
                continue
            target.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(info) as src, target.open("wb") as out:
                shutil.copyfileobj(src, out, length=1024 * 64)
    return destination


def _normalize_extracted_text(text: str) -> str:
    return "\n\n".join(part.strip() for part in text.splitlines() if part.strip())


def _truncate_chars(text: str, limit: int = _MAX_DOCUMENT_CHARS) -> str:
    cap = limit
    if _in_light_document_mode() and _SUGGEST_MAX_DOCUMENT_CHARS > 0:
        cap = min(cap, _SUGGEST_MAX_DOCUMENT_CHARS) if cap > 0 else _SUGGEST_MAX_DOCUMENT_CHARS
    if cap <= 0 or len(text) <= cap:
        return text
    return text[:cap]


def _pdf_native_text(
    path: Path,
    *,
    max_pages: int = _DEFAULT_NATIVE_PDF_MAX_PAGES,
    max_chars: int = _MAX_DOCUMENT_CHARS,
) -> tuple[str, int]:
    from pypdf import PdfReader

    light = _in_light_document_mode()
    if light:
        if _SUGGEST_OCR_MAX_PAGES > 0:
            max_pages = (
                min(max_pages, _SUGGEST_OCR_MAX_PAGES) if max_pages > 0 else _SUGGEST_OCR_MAX_PAGES
            )
        if _SUGGEST_MAX_DOCUMENT_CHARS > 0:
            max_chars = (
                min(max_chars, _SUGGEST_MAX_DOCUMENT_CHARS)
                if max_chars > 0
                else _SUGGEST_MAX_DOCUMENT_CHARS
            )
    reader = PdfReader(path)
    total_pages = len(reader.pages)
    if total_pages <= 0:
        return "", 0
    if max_pages <= 0:
        page_indexes = list(range(total_pages))
    elif light:
        # Random interior pages for profile/persona skims (1-based → 0-based).
        page_indexes = [n - 1 for n in _sample_page_numbers(total_pages, max_pages)]
    else:
        page_indexes = list(range(min(total_pages, max_pages)))
    parts: list[str] = []
    used = 0
    for index in page_indexes:
        raw = (reader.pages[index].extract_text() or "").strip()
        normalized = _normalize_extracted_text(raw)
        if not normalized:
            continue
        if max_chars > 0 and used + len(normalized) > max_chars:
            remaining = max_chars - used
            if remaining > 0:
                parts.append(normalized[:remaining])
            break
        parts.append(normalized)
        used += len(normalized)
    # Keep page boundaries for chunk_level=page sampling.
    return "\f".join(parts), total_pages


def _pdf_text_is_usable(text: str, page_count: int) -> bool:
    if page_count <= 0:
        return bool(text.strip())
    parts = [part for part in text.split("\f") if part.strip()] if "\f" in text else [part for part in text.split("\n\n") if part.strip()]
    nonempty_pages = len(parts)
    chars_per_page = len(text.replace("\f", "")) / page_count
    nonempty_ratio = nonempty_pages / page_count
    return chars_per_page >= _MIN_PDF_CHARS_PER_PAGE and nonempty_ratio >= _MIN_PDF_NONEMPTY_PAGE_RATIO


def _pdf_page_count(path: Path) -> int:
    from pypdf import PdfReader

    return len(PdfReader(path).pages)


def _configure_ocr_runtime() -> None:
    """Pin OpenMP threads so parallel page workers do not oversubscribe CPUs."""
    os.environ["OMP_THREAD_LIMIT"] = str(_OCR_OMP_THREADS)


def _ocr_one_page(
    path: Path,
    page_num: int,
    *,
    dpi: int,
) -> tuple[int, str]:
    """Render + OCR a single page; release the image before returning."""
    import pytesseract
    from pdf2image import convert_from_path

    _configure_ocr_runtime()
    images = convert_from_path(
        str(path),
        dpi=dpi,
        first_page=page_num,
        last_page=page_num,
        grayscale=True,
        thread_count=1,
    )
    if not images:
        return page_num, ""
    image = images[0]
    try:
        raw = pytesseract.image_to_string(image)
    finally:
        image.close()
        del images
    return page_num, _normalize_extracted_text(raw)


def _ocr_pdf(
    path: Path,
    *,
    dpi: int = _DEFAULT_OCR_DPI,
    max_pages: int = _DEFAULT_OCR_MAX_PAGES,
    workers: int = _DEFAULT_OCR_WORKERS,
) -> tuple[str, int, int]:
    """OCR pages. Returns (text, pages_attempted, total_pages)."""
    if _in_light_document_mode():
        dpi = min(dpi, _SUGGEST_OCR_DPI)
        if _SUGGEST_OCR_MAX_PAGES > 0:
            max_pages = (
                min(max_pages, _SUGGEST_OCR_MAX_PAGES) if max_pages > 0 else _SUGGEST_OCR_MAX_PAGES
            )
    try:
        import pytesseract
        from pdf2image import convert_from_path  # noqa: F401 — import check
        from pdf2image.exceptions import PDFInfoNotInstalledError, PDFPageCountError
    except ImportError as exc:
        raise RuntimeError(
            "OCR fallback requires optional deps: pip install 'ds-structure-server[ocr]' "
            f"(pytesseract + pdf2image). Missing while reading: {path}"
        ) from exc

    _configure_ocr_runtime()
    try:
        total_pages = _pdf_page_count(path)
    except Exception as exc:
        raise RuntimeError(f"Could not read PDF page count for OCR: {path}") from exc
    if total_pages <= 0:
        return "", 0, 0
    if max_pages <= 0:
        page_numbers = list(range(1, total_pages + 1))
    elif _in_light_document_mode():
        page_numbers = _sample_page_numbers(total_pages, max_pages)
    else:
        page_numbers = list(range(1, min(total_pages, max_pages) + 1))
    worker_count = max(1, min(workers, len(page_numbers)))
    logger.info(
        "ocr: start file={} dpi={} pages={}/{} sampled={} workers={} omp_threads={}",
        path.name,
        dpi,
        len(page_numbers),
        total_pages,
        page_numbers[:8] + (["…"] if len(page_numbers) > 8 else []),
        worker_count,
        _OCR_OMP_THREADS,
    )

    results: dict[int, str] = {}
    done = 0

    def _run(page_num: int) -> tuple[int, str]:
        try:
            return _ocr_one_page(path, page_num, dpi=dpi)
        except pytesseract.TesseractNotFoundError as exc:
            raise RuntimeError(
                f"OCR fallback requires the `tesseract` binary on PATH: {path}"
            ) from exc
        except (PDFInfoNotInstalledError, PDFPageCountError, OSError) as exc:
            raise RuntimeError(
                f"OCR fallback requires poppler (`pdftoppm`) to render PDF pages: {path}"
            ) from exc

    try:
        if worker_count == 1:
            for page_num in page_numbers:
                _, text = _run(page_num)
                if text.strip():
                    results[page_num] = text
                done += 1
                if done == 1 or done % 10 == 0 or done == len(page_numbers):
                    logger.info("ocr: page {}/{} nonempty={}", done, len(page_numbers), len(results))
        else:
            with ThreadPoolExecutor(max_workers=worker_count) as pool:
                futures = {pool.submit(_run, page_num): page_num for page_num in page_numbers}
                for future in as_completed(futures):
                    page_num, text = future.result()
                    if text.strip():
                        results[page_num] = text
                    done += 1
                    if done == 1 or done % 10 == 0 or done == len(page_numbers):
                        logger.info(
                            "ocr: progress {}/{} nonempty={}",
                            done,
                            len(page_numbers),
                            len(results),
                        )
    except MemoryError:
        logger.warning("ocr: stopped early on MemoryError after {} pages", done)

    ordered = [results[i] for i in sorted(results)]
    logger.info("ocr: finished pages_attempted={} nonempty={}", done, len(ordered))
    return "\f".join(ordered), done, total_pages


def _pdf_capacity_chars(sample_text: str, *, pages_attempted: int, total_pages: int) -> int:
    """Extrapolate corpus size from an OCR skim so suggest row counts stay full-document."""
    sample_chars = len((sample_text or "").replace("\f", ""))
    if sample_chars <= 0:
        return 0
    if pages_attempted <= 0 or total_pages <= pages_attempted:
        return sample_chars
    return max(sample_chars, int(sample_chars * (total_pages / pages_attempted)))


def _read_pdf(path: Path) -> tuple[str, dict[str, Any]]:
    logger.info("pdf: native extract {}", path.name)
    text, page_count = _pdf_native_text(path)
    if _pdf_text_is_usable(text, page_count):
        chars = len(text.replace("\f", ""))
        logger.info("pdf: native usable pages={} chars={}", page_count, chars)
        return text, {"page_count": page_count, "capacity_chars": chars}
    logger.info("pdf: native sparse pages={} chars={} → OCR fallback", page_count, len(text))
    ocr_text, pages_attempted, total_pages = _ocr_pdf(path)
    if ocr_text.strip():
        capacity_chars = _pdf_capacity_chars(
            ocr_text, pages_attempted=pages_attempted, total_pages=total_pages
        )
        logger.info(
            "pdf: OCR ok sample_chars={} capacity_chars={} pages={}/{}",
            len(ocr_text.replace("\f", "")),
            capacity_chars,
            pages_attempted,
            total_pages,
        )
        return ocr_text, {
            "page_count": total_pages,
            "capacity_chars": capacity_chars,
            "ocr_pages": pages_attempted,
        }
    if text.strip():
        logger.warning("pdf: OCR empty; using sparse native text chars={}", len(text))
        return text, {"page_count": page_count, "capacity_chars": len(text.replace("\f", ""))}
    raise ValueError(f"Could not extract text from PDF via native parsing or OCR: {path}")


def _read_document(path: Path) -> tuple[str, dict[str, Any]]:
    """Return (text, extract_metadata). capacity_chars is full-doc estimate when OCR is skimmed."""
    suffix = path.suffix.casefold()
    if suffix == ".pdf":
        return _read_pdf(path)
    char_cap = _MAX_DOCUMENT_CHARS
    if _in_light_document_mode() and _SUGGEST_MAX_DOCUMENT_CHARS > 0:
        char_cap = (
            min(char_cap, _SUGGEST_MAX_DOCUMENT_CHARS) if char_cap > 0 else _SUGGEST_MAX_DOCUMENT_CHARS
        )
    if suffix == ".docx":
        from docx import Document

        parts: list[str] = []
        used = 0
        for paragraph in Document(path).paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            if char_cap > 0 and used + len(text) + 2 > char_cap:
                remaining = char_cap - used
                if remaining > 0:
                    parts.append(text[:remaining])
                break
            parts.append(text)
            used += len(text) + 2
        joined = "\n\n".join(parts)
        return joined, {"capacity_chars": len(joined)}
    if char_cap <= 0:
        text = path.read_text(encoding="utf-8", errors="replace")
        return text, {"capacity_chars": len(text)}
    chunks: list[str] = []
    used = 0
    with path.open("r", encoding="utf-8", errors="replace") as handle:
        while used < char_cap:
            block = handle.read(min(64 * 1024, char_cap - used))
            if not block:
                break
            chunks.append(block)
            used += len(block)
    text = "".join(chunks)
    return text, {"capacity_chars": len(text)}


def iter_structured_rows(path: Path, *, max_rows: int | None = None) -> Iterator[dict[str, Any]]:
    """Yield structured rows without requiring the whole file in memory (jsonl/csv)."""
    suffix = path.suffix.casefold()
    yielded = 0
    if suffix == ".jsonl":
        with path.open("r", encoding="utf-8", errors="replace") as handle:
            for line in handle:
                if max_rows is not None and yielded >= max_rows:
                    return
                text = line.strip()
                if not text:
                    continue
                value = json.loads(text)
                if isinstance(value, dict):
                    yield value
                    yielded += 1
        return
    if suffix == ".csv":
        with path.open(newline="", encoding="utf-8", errors="replace") as handle:
            for row in DictReader(handle):
                if max_rows is not None and yielded >= max_rows:
                    return
                if isinstance(row, dict):
                    yield row
                    yielded += 1
        return
    # JSON arrays/objects: size-capped full parse (true streaming needs ijson).
    if path.stat().st_size > _MAX_STRUCTURED_FILE_BYTES:
        raise ValueError(
            f"Structured JSON source exceeds {_MAX_STRUCTURED_FILE_BYTES // (1024 * 1024)}MB: {path}"
        )
    value = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    rows: list[dict[str, Any]]
    if isinstance(value, list):
        rows = [row for row in value if isinstance(row, dict)]
    elif isinstance(value, dict):
        rows = []
        for key in ("rows", "records", "documents", "data", "items", "results"):
            nested = value.get(key)
            if isinstance(nested, list) and all(isinstance(row, dict) for row in nested):
                rows = nested
                break
        if not rows:
            rows = [value]
    else:
        raise ValueError(f"Structured source must contain objects: {path}")
    for row in rows:
        if max_rows is not None and yielded >= max_rows:
            return
        yield row
        yielded += 1


def sample_and_count_structured(
    path: Path,
    *,
    sample_limit: int = 200,
) -> tuple[list[dict[str, Any]], int]:
    """Stream structured sources: keep a small sample and an exact/cheap row count."""
    suffix = path.suffix.casefold()
    sample: list[dict[str, Any]] = []
    count = 0
    if suffix in {".jsonl", ".csv"}:
        for row in iter_structured_rows(path):
            count += 1
            if len(sample) < sample_limit:
                sample.append(row)
        return sample, count
    # JSON: one bounded parse; sample + full count from the in-memory list once.
    rows = list(iter_structured_rows(path))
    return rows[:sample_limit], len(rows)


def _structured_rows(path: Path, *, max_rows: int | None = None) -> list[dict[str, Any]]:
    return list(iter_structured_rows(path, max_rows=max_rows))


def load_local_documents(
    input_path: Path,
    *,
    limit: int,
    seed: int = 42,
    schema_sample_rows: int = 32,
    text_column: str | None = None,
    id_column: str | None = None,
    title_column: str | None = None,
    metadata_columns: tuple[str, ...] = (),
) -> tuple[list[SourceDocument], DiscoveredSchema]:
    path = input_path.expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(path)

    if path.is_file() and path.suffix.casefold() in _ARCHIVE_SUFFIXES:
        unzipped = path.parent / f"{path.stem}_unzipped"
        if unzipped.exists():
            import shutil

            shutil.rmtree(unzipped, ignore_errors=True)
        path = _safe_extract_zip(path, unzipped)

    if path.is_dir():
        files = sorted(
            child
            for child in path.rglob("*")
            if child.is_file()
            and child.suffix.casefold() in (_DOCUMENT_SUFFIXES | _STRUCTURED_SUFFIXES)
        )
        random.Random(seed).shuffle(files)
        documents: list[SourceDocument] = []
        for file in files:
            if len(documents) >= limit:
                break
            suffix = file.suffix.casefold()
            if suffix in _DOCUMENT_SUFFIXES:
                text, extract_meta = _read_document(file)
                text = _truncate_chars(text)
                if text.strip():
                    meta = {"path": str(file), **extract_meta}
                    # Truncation shrinks text; keep capacity_chars as the pre-truncate estimate.
                    if "capacity_chars" not in extract_meta:
                        meta["capacity_chars"] = len(text.replace("\f", ""))
                    documents.append(
                        SourceDocument(
                            doc_id=str(file.relative_to(path)),
                            text=text,
                            title=file.stem,
                            metadata=meta,
                        )
                    )
                continue
            remaining = limit - len(documents)
            if remaining <= 0:
                break
            nested_docs, _ = load_local_documents(
                file,
                limit=remaining,
                seed=seed,
                schema_sample_rows=schema_sample_rows,
                text_column=text_column,
                id_column=id_column,
                title_column=title_column,
                metadata_columns=metadata_columns,
            )
            documents.extend(nested_docs)
        schema = DiscoveredSchema(
            source=str(path),
            source_kind="document-directory",
            text_column="$document",
            id_column="$relative_path",
            title_column="$filename",
            metadata_columns=("path",),
            sampled_rows=min(len(files), limit),
            confidence=1.0,
        )
        return [document for document in documents if document.text.strip()], schema

    suffix = path.suffix.casefold()
    if suffix in _DOCUMENT_SUFFIXES:
        text, extract_meta = _read_document(path)
        text = _truncate_chars(text)
        meta = {"path": str(path), **extract_meta}
        if "capacity_chars" not in extract_meta:
            meta["capacity_chars"] = len(text.replace("\f", ""))
        schema = DiscoveredSchema(
            source=str(path),
            source_kind="document",
            text_column="$document",
            id_column="$path",
            title_column="$filename",
            metadata_columns=("path",),
            sampled_rows=1,
            confidence=1.0,
        )
        return [SourceDocument(str(path), text, path.stem, meta)], schema
    if suffix not in _STRUCTURED_SUFFIXES:
        raise ValueError(f"Unsupported local source type: {suffix or '<none>'}")

    rows = _structured_rows(path, max_rows=max(limit, schema_sample_rows))
    sample = rows[:schema_sample_rows]
    schema = discover_schema(
        sample,
        source=str(path),
        source_kind=suffix.removeprefix("."),
        text_column=text_column,
        id_column=id_column,
        title_column=title_column,
        metadata_columns=metadata_columns,
    )
    random.Random(seed).shuffle(rows)
    documents = list(
        documents_from_rows(
            rows,
            source=str(path),
            schema=schema,
            limit=limit,
        )
    )
    return documents, schema


def _fallback_id(dataset_id: str, index: int, text: str) -> str:
    digest = hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]
    return f"{dataset_id}:{index}:{digest}"


def documents_from_rows(
    rows: Iterator[dict[str, Any]] | list[dict[str, Any]],
    *,
    source: str,
    schema: DiscoveredSchema,
    limit: int,
) -> Iterator[SourceDocument]:
    emitted = 0
    for index, row in enumerate(rows):
        raw_text = _joined_text(row, schema.text_column)
        if not raw_text:
            continue
        raw_id = nested_value(row, schema.id_column)
        doc_id = str(raw_id).strip() if raw_id not in (None, "") else _fallback_id(source, index, raw_text)
        raw_title = nested_value(row, schema.title_column)
        title = str(raw_title).strip() if raw_title not in (None, "") else doc_id
        metadata = {column: nested_value(row, column) for column in schema.metadata_columns}
        yield SourceDocument(doc_id=doc_id, text=raw_text, title=title, metadata=metadata)
        emitted += 1
        if emitted >= limit:
            return


def stream_documents(
    dataset_id: str,
    text_column: str,
    *,
    split: str = "train",
    config: str | None = None,
    id_column: str | None = None,
    title_column: str | None = None,
    metadata_columns: tuple[str, ...] = (),
    limit: int = 3,
    seed: int = 42,
    shuffle_buffer: int = 1_000,
) -> Iterator[SourceDocument]:
    dataset = load_dataset(dataset_id, config, split=split, streaming=True)
    if shuffle_buffer > 1:
        dataset = dataset.shuffle(seed=seed, buffer_size=shuffle_buffer)
    schema = DiscoveredSchema(
        source=dataset_id,
        source_kind="huggingface",
        text_column=text_column,
        id_column=id_column,
        title_column=title_column,
        metadata_columns=metadata_columns,
        sampled_rows=0,
        confidence=1.0,
    )
    iterator = iter(dataset)
    try:
        yield from documents_from_rows(iterator, source=dataset_id, schema=schema, limit=limit)
    finally:
        close = getattr(iterator, "close", None)
        if callable(close):
            close()
        del iterator
        del dataset
        gc.collect()
